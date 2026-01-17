from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=0, bottom=0, start=0, end=0):
    """
    Helper function to set cell margins tightly for the layout.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for side, value in [('top', top), ('bottom', bottom), ('left', start), ('right', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(int(value * 1440))) # 1440 twips = 1 inch
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)

def create_serif_resume():
    doc = Document()

    # --- 1. SETUP: Margins & Fonts ---
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # Define the "Blue" color
    accent_color = RGBColor(46, 85, 160)

    # --- 2. HEADER SECTION ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(7.5)
    header_table.autofit = False
    
    # Name
    cell_h1 = header_table.cell(0, 0)
    cell_h1.width = Inches(4.5)
    name_para = cell_h1.paragraphs[0]
    name_run = name_para.add_run("Jamil Ahmad Rupak")
    name_run.bold = True
    name_run.font.size = Pt(26)
    name_run.font.name = 'Times New Roman'
    
    title_para = cell_h1.add_paragraph()
    title_run = title_para.add_run("Software Developer")
    title_run.font.size = Pt(14)
    title_run.font.name = 'Times New Roman'
    title_run.italic = True
    
    # Contact
    cell_h2 = header_table.cell(0, 1)
    cell_h2.width = Inches(3.0)
    contact_text = (
        "Dhaka, Bangladesh\n"
        "01771823979\n"
        "jhrupok@gmail.com\n"
        "linkedin.com/in/jamil-ahmad-rupak"
    )
    contact_para = cell_h2.paragraphs[0]
    contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(9)
    contact_run.font.name = 'Times New Roman'

    doc.add_paragraph() # Spacer

    # --- 3. MAIN LAYOUT (2 Columns) ---
    layout_table = doc.add_table(rows=1, cols=2)
    layout_table.autofit = False

    # Define Widths
    col1 = layout_table.cell(0, 0)
    col1.width = Inches(4.8)
    set_cell_margins(col1, end=0.3) 
    
    col2 = layout_table.cell(0, 1)
    col2.width = Inches(2.7)
    
    # Helper to add Section Headers
    def add_section_header(cell, text):
        p = cell.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.color.rgb = accent_color
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    # === LEFT COLUMN ===
    add_section_header(col1, "Profile")
    summary = (
        "A driven and adaptable software developer with knowledge of frontend and backend programming. "
        "Competent in leveraging frameworks like React, Next.js, and Django to secure online apps, build REST APIs, "
        "and establish database interfaces. Skilled in Python, Java, and JavaScript, with practical knowledge of MySQL "
        "and PostgreSQL."
    )
    col1.add_paragraph(summary).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_section_header(col1, "Projects")
    
    # Proj 1
    p1 = col1.add_paragraph()
    p1.add_run("ModelMate | AI-User’s Community").bold = True
    p1.add_run("\nDjango, PostgreSQL, REST APIs").italic = True
    
    # Use standard bullets here (they look fine in wide columns)
    ul1 = col1.add_paragraph(style='List Bullet')
    ul1.add_run("Developed a dynamic community platform enabling interaction among AI users.")
    ul1.paragraph_format.space_after = Pt(0)
    ul2 = col1.add_paragraph(style='List Bullet')
    ul2.add_run("Implemented secure REST APIs to handle user data and established efficient database interfaces.")

    # Proj 2
    p2 = col1.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    p2.add_run("LTSR | Productivity App").bold = True
    p2.add_run("\nReact, Next.js").italic = True
    
    ul3 = col1.add_paragraph(style='List Bullet')
    ul3.add_run("Designed a productivity application with a unique concept to enhance user efficiency.")
    ul3.paragraph_format.space_after = Pt(0)
    ul4 = col1.add_paragraph(style='List Bullet')
    ul4.add_run("Utilized React/Next.js to build a responsive, scalable frontend interface.")

    add_section_header(col1, "Education")
    edu1 = col1.add_paragraph()
    edu1.add_run("B.Sc. in Computer Science and Engineering").bold = True
    edu1.add_run("\nChittagong University of Engineering & Technology")
    edu1.add_run("\nSession: 2021-22")
    
    edu2 = col1.add_paragraph()
    edu2.paragraph_format.space_before = Pt(8)
    edu2.add_run("Higher Secondary Certificate (HSC)").bold = True
    edu2.add_run("\nBAF Shaheen College, Dhaka")
    edu2.add_run("\nGPA: 5.00 | Year: 2021")

    # === RIGHT COLUMN (Fixing the dots) ===
    
    # Helper for manual, tight bullets
    def add_tight_bullet(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(2)
        # Using a manual bullet character (u2022) + tab
        run = p.add_run(f"\u2022\t{text}") 
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    add_section_header(col2, "Skills")
    skills_list = [
        ("Languages", "Python, Java, C++, C#, JS"),
        ("Frameworks", "Django, React, Next.js"),
        ("Databases", "PostgreSQL, MySQL"),
        ("Concepts", "OOP, Data Structures, REST"),
        ("Tools", "Git, Spring Security")
    ]
    for cat, items in skills_list:
        p = col2.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{cat}:").bold = True
        p.add_run(f"\n{items}")

    add_section_header(col2, "Awards")
    awards = [
        "President Scouts Award",
        "Shapla Cub Award",
        "Former Secretary, Math Club"
    ]
    for award in awards:
        add_tight_bullet(col2, award)

    add_section_header(col2, "Interests")
    interests = ["Traditional Food", "Cricket", "Complex Problem Solving"]
    for interest in interests:
        add_tight_bullet(col2, interest)

    add_section_header(col2, "Links")
    links = col2.add_paragraph()
    links.add_run("Portfolio: ").bold = True
    links.add_run("jamilahmadrupak.vercel.app\n")
    links.add_run("GitHub: ").bold = True
    links.add_run("github.com/JamilAhmadRupak")

    doc.save('Jamil_Ahmad_Rupak_Serif_Resume.docx')

if __name__ == "__main__":
    create_serif_resume()